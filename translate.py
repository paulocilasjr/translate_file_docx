import os
import subprocess
import tempfile
import pymupdf
from docx import Document
from deep_translator import GoogleTranslator
import csv
import argparse
import zipfile
from PIL import Image, ImageDraw, ImageFont
import pytesseract
import io

def chunk_text(text, max_len=4000):
    """Split text into chunks for translation."""
    chunks = []
    current = ''
    for word in text.split():
        if len(current) + len(word) + 1 > max_len:
            chunks.append(current.strip())
            current = word + ' '
        else:
            current += word + ' '
    if current:
        chunks.append(current.strip())
    return chunks

def translate_text(text):
    """Translate text to Portuguese, handling large texts by chunking."""
    if not text.strip():
        return text
    translator = GoogleTranslator(source='auto', target='pt')
    if len(text) > 4000:
        chunks = chunk_text(text)
        translated_chunks = [translator.translate(chunk) for chunk in chunks]
        return ' '.join(translated_chunks)
    else:
        return translator.translate(text)

def translate_docx(input_path, output_path):
    """Translate DOCX file, including text in embedded images."""
    try:
        # First, translate regular text
        doc = Document(input_path)
        for paragraph in doc.paragraphs:
            paragraph.text = translate_text(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.text = translate_text(paragraph.text)

        # Now, handle images: Extract, OCR, translate, overlay, replace
        with tempfile.TemporaryDirectory() as temp_dir:
            # Unzip DOCX
            with zipfile.ZipFile(input_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)

            media_dir = os.path.join(temp_dir, 'word', 'media')
            if os.path.exists(media_dir):
                for img_file in os.listdir(media_dir):
                    if img_file.lower().endswith(('.png', '.jpg', '.jpeg')):
                        img_path = os.path.join(media_dir, img_file)
                        # OCR the image
                        img_text = pytesseract.image_to_string(Image.open(img_path))
                        if img_text.strip():
                            translated_text = translate_text(img_text)
                            # Overlay translated text on image (simple bottom overlay; adjust as needed)
                            img = Image.open(img_path)
                            draw = ImageDraw.Draw(img)
                            font = ImageFont.load_default()  # Use better font if available
                            draw.rectangle(((0, img.height - 30), (img.width, img.height)), fill="white")
                            draw.text((10, img.height - 30), translated_text, fill="black", font=font)
                            img.save(img_path)

            # Re-zip to new DOCX
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        zip_out.write(os.path.join(root, file), os.path.relpath(os.path.join(root, file), temp_dir))

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
    except Exception as e:
        print(f"Failed to translate DOCX {input_path}: {e}. Skipping.")

def translate_pdf(input_path, output_path):
    """Translate PDF using PyMuPDF, preserving formatting."""
    try:
        WHITE = pymupdf.pdfcolor["white"]
        doc = pymupdf.open(input_path)
        to_pt = GoogleTranslator(source='auto', target='pt')
        ocg_xref = doc.add_ocg("Portuguese", on=True)
        for page in doc:
            blocks = page.get_text("blocks", flags=pymupdf.TEXT_DEHYPHENATE)
            for block in blocks:
                bbox = pymupdf.Rect(block[:4])
                original_text = block[4]
                if not isinstance(original_text, str):
                    continue  # Skip non-string blocks
                translated = translate_text(original_text)
                if translated is None:
                    translated = ""
                page.draw_rect(bbox, color=None, fill=WHITE, oc=ocg_xref)
                page.insert_htmlbox(bbox, translated, oc=ocg_xref)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.subset_fonts()
        doc.ez_save(output_path)
    except Exception as e:
        print(f"Failed to translate PDF {input_path}: {e}. Skipping.")

def convert_using_pandoc(input_path, output_dir, to_format):
    """Convert file using Pandoc."""
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    output_path = os.path.join(output_dir, f"{base_name}.{to_format}")
    cmd = ['pandoc', input_path, '-o', output_path]
    if to_format == 'rtf':
        cmd.insert(-1, '--to=rtf')  # Specify format for RTF
    try:
        subprocess.run(cmd, check=True)
    except FileNotFoundError:
        raise FileNotFoundError("Pandoc not found. Install it via 'brew install pandoc' on macOS.")
    except Exception as e:
        print(f"Pandoc error for {input_path}: {e}")
        raise e

def process_file(input_path, output_path, ext):
    """Process file based on extension."""
    if ext == '.docx':
        translate_docx(input_path, output_path)
    elif ext == '.pdf':
        translate_pdf(input_path, output_path)
    else:  # .doc or .rtf
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                # Convert to docx
                convert_using_pandoc(input_path, temp_dir, 'docx')
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                docx_path = os.path.join(temp_dir, f"{base_name}.docx")
                # Translate the docx (including images)
                translate_docx(docx_path, docx_path)  # Overwrite
                # Convert back to original format
                original_format = ext[1:]  # e.g., 'doc' or 'rtf'
                temp_output_dir = os.path.dirname(output_path)
                os.makedirs(temp_output_dir, exist_ok=True)
                convert_using_pandoc(docx_path, temp_output_dir, original_format)
                # The converted file is {base_name}.{original_format} in temp_output_dir
                converted_path = os.path.join(temp_output_dir, f"{base_name}.{original_format}")
                if converted_path != output_path:
                    os.rename(converted_path, output_path)
        except Exception as e:
            print(f"Failed to process {ext.upper()} file {input_path}: {e}. Skipping. Install Pandoc if not already.")

def main():
    parser = argparse.ArgumentParser(description="Translate files to Portuguese.")
    parser.add_argument('--csv', type=str, help='Path to CSV file with specific files to translate')
    args = parser.parse_args()

    input_dir = "translate"
    output_root = "Portuguese"

    if args.csv:
        with open(args.csv, 'r') as f:
            reader = csv.reader(f)
            next(reader)  # Skip header 'Path'
            for row in reader:
                if row:
                    input_path = row[0]
                    if not os.path.exists(input_path):
                        print(f"File not found: {input_path}. Skipping.")
                        continue
                    file = os.path.basename(input_path)
                    if file.startswith('.'):
                        print(f"Skipping hidden file: {file}")
                        continue
                    ext = os.path.splitext(file)[1].lower()
                    if ext in ['.docx', '.doc', '.pdf', '.rtf']:
                        rel_path = os.path.relpath(os.path.dirname(input_path), input_dir)
                        output_dir = os.path.join(output_root, rel_path)
                        output_path = os.path.join(output_dir, file)
                        print(f"Translating {input_path} to {output_path}")
                        process_file(input_path, output_path, ext)
    else:
        for root, dirs, files in os.walk(input_dir):
            for file in files:
                if file.startswith('.'):
                    print(f"Skipping hidden file: {file}")
                    continue
                ext = os.path.splitext(file)[1].lower()
                if ext in ['.docx', '.doc', '.pdf', '.rtf']:
                    input_path = os.path.join(root, file)
                    rel_path = os.path.relpath(root, input_dir)
                    output_dir = os.path.join(output_root, rel_path)
                    output_path = os.path.join(output_dir, file)
                    print(f"Translating {input_path} to {output_path}")
                    process_file(input_path, output_path, ext)

if __name__ == "__main__":
    main()
